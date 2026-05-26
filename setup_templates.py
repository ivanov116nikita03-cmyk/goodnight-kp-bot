#!/usr/bin/env python3
"""
cd /opt/gnbot && python3 setup_templates.py
"""
import os, sys, re, zipfile, io

try:
    from docx import Document
    from docx.text.paragraph import Paragraph as DocP
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip3 install python-docx")
    from docx import Document
    from docx.text.paragraph import Paragraph as DocP
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def replace_all(body, pairs):
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para = DocP(child, body)
            full = ''.join(r.text for r in para.runs)
            nw = full
            for old, new in pairs:
                if old: nw = nw.replace(old, new)
            if nw != full and para.runs:
                para.runs[0].text = nw
                for r in para.runs[1:]: r.text = ''
        elif tag in ('tbl','tr','tc','body'):
            replace_all(child, pairs)


def replace_para_regex(body, pattern, repl):
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para = DocP(child, body)
            full = ''.join(r.text for r in para.runs)
            nw = re.sub(pattern, repl, full)
            if nw != full and para.runs:
                para.runs[0].text = nw
                for r in para.runs[1:]: r.text = ''
        elif tag in ('tbl','tr','tc','body'):
            replace_para_regex(child, pattern, repl)


def fix_standalone(body, find_text, replace_text):
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para = DocP(child, body)
            if para.text.strip() == find_text:
                if para.runs:
                    para.runs[0].text = replace_text
                    for r in para.runs[1:]: r.text = ''
        elif tag in ('tbl','tr','tc','body'):
            fix_standalone(child, find_text, replace_text)


def fix_docx_formatting(path):
    """Применяет к готовому .docx: шрифт 10pt, убирает жёлтый, 
    исправляет выравнивание секций 4 и 6 (не трогая заголовок документа)."""
    
    # Читаем все файлы архива
    files = {}
    with zipfile.ZipFile(path, 'r') as z:
        for name in z.namelist():
            files[name] = z.read(name)

    xml = files['word/document.xml'].decode('utf-8')

    # 1. Убираем жёлтую подсветку
    xml = re.sub(r'<w:highlight[^/]*/>', '', xml)

    # 2. Шрифт 10pt для всего текста
    xml = re.sub(r'<w:sz w:val="\d+"/>', '<w:sz w:val="20"/>', xml)
    xml = re.sub(r'<w:szCs w:val="\d+"/>', '<w:szCs w:val="20"/>', xml)

    # 3. Выравнивание: убираем center у параграфов КОТОРЫЕ НЕ являются
    #    заголовком документа. Заголовок ("Договор" / "на оказание услуг")
    #    получает центрирование через стиль — inline <w:jc> им не нужен.
    xml = xml.replace('<w:jc w:val="center"/>', '<w:jc w:val="left"/>')

    files['word/document.xml'] = xml.encode('utf-8')

    # Пересобираем zip
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    with open(path, 'wb') as f:
        f.write(buf.getvalue())


# ════════════════════════════════════════════════════════════
# ДОГОВОР
# ════════════════════════════════════════════════════════════
print("Создаю template_dogovor_new.docx...")

# Используем template_dogovor.docx (конвертированный оригинал, не генерацию)
src = 'template_dogovor.docx'
if not os.path.exists(src):
    print("ОШИБКА: нет template_dogovor.docx")
    sys.exit(1)

doc = Document(src)

# Исправляем реквизиты исполнителя
replace_all(doc.element.body, [
    ('ИНН 032386861274', 'ИНН 032315540193'),
    ('р/с 40802810000009415686', 'р/с 40802810500007432200'),
    ('ochrvvv@mail.ru', 'msk@goodnight.show'),
])

# Номер договора (любой номер → маркер)
replace_para_regex(doc.element.body,
    r'(на оказание услуг № )\d+',
    r'\1[[НОМ]]')

# Дата договора (разбита по runs — ловим через полный текст параграфа)
for para in doc.paragraphs:
    t = ''.join(r.text for r in para.runs)
    nw = re.sub(r'«\d+»\s*[а-яёА-ЯЁ]+ \d{4}', '«[[ДЕНЬ]]» [[МЕС_ГОД]]', t)
    if nw != t and para.runs:
        para.runs[0].text = nw
        for r in para.runs[1:]: r.text = ''

# Директор заказчика в шапке
for para in doc.paragraphs:
    if 'Генерального Директора' in para.text and 'действующего' in para.text:
        full = ''.join(r.text for r in para.runs)
        nw = re.sub(
            r'Генерального Директора\s+[А-ЯЁ].+?,\s*действующего',
            'Генерального Директора [[ДИР]], действующего', full)
        if nw != full and para.runs:
            para.runs[0].text = nw
            for r in para.runs[1:]: r.text = ''
        break

# Динамические поля (длительность, срок, место, стоимость)
for para in doc.paragraphs:
    full = ''.join(r.text for r in para.runs)
    nw = full
    nw = re.sub(r'(продолжительность оказания Услуг\s*[–\-]\s*)[\d,\.]+\s*час[а-я]*;',
                r'\1[[ДЛИТ]];', nw)
    nw = re.sub(r'(Срок оказания Услуг\s*[–\-]\s*)\d+\s+[а-яё]+\s+\d{4}\s+года\s+[\d:]+\s*[-–]\s*[\d:]+;',
                r'\1[[ДАТА_МЕР]];', nw)
    nw = re.sub(r'(Место проведения:\s*).+', r'\1[[МЕСТО]]', nw)
    nw = re.sub(r'(\bсоставляет\s+)[\d\s\xa0]+\([^)]+\)(\s+рублей)',
                r'\1[[СУМ_Ц]] ([[СУМ_СЛ]])\2', nw)
    if nw != full and para.runs:
        para.runs[0].text = nw
        for r in para.runs[1:]: r.text = ''

# Реквизиты заказчика (блок 7) — маркеры
replace_all(doc.element.body, [
    ('ОГРН ', 'ОГРН [[ОГРН]]'),
    ('ИНН / КПП /', 'ИНН [[ИНН]] / КПП [[КПП]]'),
    ('Расчетный счет №', 'Расчетный счет № [[РС]]'),
    ('к/с №', 'к/с № [[КС]]'),
])
replace_para_regex(doc.element.body,
    r'(Юридический адрес:\s*).+', r'\1[[ЗАК_АДР]]')
replace_para_regex(doc.element.body,
    r'(Фактический адрес:\s*).+', r'\1[[ЗАК_АДР]]')
fix_standalone(doc.element.body, 'БИК', 'БИК [[БИК]]')
fix_standalone(doc.element.body, 'Заказчик', '[[ЗАК]]')

# Также заменяем полное название компании-заказчика (шапка и блок 7)
replace_para_regex(doc.element.body,
    r'(Общество с ограниченной ответственностью|ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ)\s*[«""][^»""]+[»""]',
    '[[ЗАК]]')

# Подписи (блок 8)
replace_all(doc.element.body, [
    ('Генеральный директор Заказчик', 'Генеральный директор [[ЗАК]]'),
])
for para in doc.paragraphs:
    t = para.text.strip()
    # Строка только из ФИО → маркер директора
    if re.match(r'^[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+$', t):
        if para.runs:
            para.runs[0].text = '[[ДИР]]'
            for r in para.runs[1:]: r.text = ''
    # Инициалы в подписи
    if para.runs:
        full = ''.join(r.text for r in para.runs)
        nw = re.sub(r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.)', '[[ДИР_ИНИ]]', full)
        if nw != full:
            para.runs[0].text = nw
            for r in para.runs[1:]: r.text = ''

doc.save('template_dogovor_new.docx')

# Применяем форматирование: 10pt, убираем жёлтый, центр→лево
fix_docx_formatting('template_dogovor_new.docx')
sz = os.path.getsize('template_dogovor_new.docx')
print(f"  template_dogovor_new.docx: {sz} байт — OK")


# ════════════════════════════════════════════════════════════
# СЧЁТ
# ════════════════════════════════════════════════════════════
print("Создаю template_schet_new.docx...")

if not os.path.exists('template_schet.docx'):
    print("ОШИБКА: нет template_schet.docx")
    sys.exit(1)

doc2 = Document('template_schet.docx')
replace_all(doc2.element.body, [
    ('ИП ЭРДЫНЕЕВ ГЭСЭР БУЯНТУЕВИЧ', 'ИП Очирова Оксана Эдуардовна'),
    ('670011, РОССИЯ, РЕСП БУРЯТИЯ, Г УЛАН-УДЭ, МКР 142-Й, -, Д 4, КВ 18',
     '670031, РОССИЯ, РЕСП БУРЯТИЯ, Г УЛАН-УДЭ, ПР СТРОИТЕЛЕЙ, Д 62, КВ 49'),
    ('ЮЛ, ИНН, КПП, АДРЕС', '[[ЗАК_ПОЛН]]'),
    ('Счет на оплату № 304', 'Счет на оплату № [[НОМ]]'),
    ('от 10 сентября 2025 г.', 'от [[ДАТА_СЧЕТ]] г.'),
    ('Организация мероприятия 10.10.2025', 'Организация мероприятия [[ДАТА_МЕР]]'),
    ('1 000,00', '[[СУМ_Ц]]'),
    ('1\xa0000,00', '[[СУМ_Ц]]'),
    ('1000,00 руб.', '[[СУМ_Ц]] руб.'),
    ('сумму 1000,00', 'сумму [[СУМ_Ц]]'),
    ('Одна тысяча рублей 00 копеек', '[[СУМ_СЛ]]'),
])
doc2.save('template_schet_new.docx')
fix_docx_formatting('template_schet_new.docx')
sz2 = os.path.getsize('template_schet_new.docx')
print(f"  template_schet_new.docx: {sz2} байт — OK")

print("\nГотово! Перезапусти бот:")
print("systemctl restart gnbot && systemctl status gnbot")
